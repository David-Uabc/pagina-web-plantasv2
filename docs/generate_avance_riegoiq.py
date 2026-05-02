from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT_PATH = r"C:\Users\da_sa\OneDrive\Desktop\pagina web plantas\docs\Avance_RiegoIQ_2_semanas.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
      tc_mar = OxmlElement("w:tcMar")
      tc_pr.append(tc_mar)
    for margin_name, margin_value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
      node = tc_mar.find(qn(f"w:{margin_name}"))
      if node is None:
        node = OxmlElement(f"w:{margin_name}")
        tc_mar.append(node)
      node.set(qn("w:w"), str(margin_value))
      node.set(qn("w:type"), "dxa")


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text)


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    title = styles["Title"]
    title.font.name = "Arial"
    title.font.size = Pt(22)
    title.font.bold = True
    title.font.color.rgb = RGBColor(17, 24, 39)

    for heading_name, size in [("Heading 1", 16), ("Heading 2", 13)]:
        style = styles[heading_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(22, 101, 52)

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = header_p.add_run("RiegoIQ | Informe de avance técnico")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 116, 139)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer_p.add_run("Proyecto de titulación | Últimas 2 semanas")
    footer_run.font.name = "Arial"
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor(100, 116, 139)


def add_cover(doc):
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.add_run("Informe de Avance Técnico\nRiegoIQ")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.LEFT
    subtitle_run = subtitle.add_run(
        "Resumen de mejoras, correcciones y preparación para producción realizadas durante las últimas dos semanas."
    )
    subtitle_run.font.name = "Arial"
    subtitle_run.font.size = Pt(12)
    subtitle_run.font.color.rgb = RGBColor(71, 85, 105)

    info = doc.add_table(rows=4, cols=2)
    info.style = "Table Grid"
    info.columns[0].width = Inches(2)
    info.columns[1].width = Inches(4.5)
    metadata = [
        ("Proyecto", "RiegoIQ - Sistema de riego IoT"),
        ("Alumno", "David Reséndiz"),
        ("Periodo", "Avance de 2 semanas"),
        ("Enfoque", "Frontend, backend, firmware, seguridad, rendimiento y despliegue en Render"),
    ]
    for row, (label, value) in zip(info.rows, metadata):
        row.cells[0].text = label
        row.cells[1].text = value
        set_cell_shading(row.cells[0], "E2F5EA")
        for cell in row.cells:
            set_cell_margins(cell)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(11)

    doc.add_paragraph("")


def add_status_table(doc):
    doc.add_heading("1. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "Durante este periodo se trabajó en estabilizar la aplicación completa, corregir errores funcionales y de despliegue, "
        "separar correctamente los datos por usuario, integrar mejor los ESP32, optimizar el rendimiento de la interfaz y "
        "preparar el sistema para pruebas reales y despliegue en Render."
    )

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["Área", "Resultado", "Impacto"]
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, "D9F2E3")
        set_cell_margins(cell)

    rows = [
        ("Frontend", "Rediseño, corrección de menús, búsqueda, accesibilidad visual y mejor comportamiento móvil/desktop", "Mayor calidad visual y mejor experiencia de usuario"),
        ("Backend", "Aislamiento por usuario, validaciones, CORS correcto, sesiones estables y rutas más seguras", "Menos errores entre cuentas y mejor consistencia del sistema"),
        ("Firmware", "Soporte completo para ambos sectores y cola offline de lecturas", "Mayor robustez en pruebas reales con ESP32"),
        ("Producción", "Preparación de variables de entorno, Render, rutas SPA y corrección de recursos 404", "Despliegue funcional en la nube"),
    ]
    for area, result, impact in rows:
        row = table.add_row().cells
        row[0].text = area
        row[1].text = result
        row[2].text = impact
        for cell in row:
            set_cell_margins(cell)


def add_bug_section(doc):
    doc.add_heading("2. Bugs y errores corregidos", level=1)
    items = [
        "Se corrigió el menú del usuario en el navbar para que volviera a desplegar perfil, configuración, apariencia y cerrar sesión.",
        "Se resolvió el problema de clics internos del menú renderizado por portal, evitando que se cerrara antes de abrir los modales.",
        "Se eliminó el warning de React por claves duplicadas dentro de AnimatePresence en Navbar.",
        "Se reparó la búsqueda global para que el overlay aparezca correctamente y no quede bugueado dentro del header.",
        "Se corrigieron problemas de codificación visual como caracteres raros, escapes unicode visibles y mojibake en múltiples componentes.",
        "Se arreglaron errores de recarga en rutas SPA como /login, /superior e /inferior mediante rewrites correctos para Render.",
        "Se corrigieron 404 del service worker, del manifest y de iconos inexistentes.",
        "Se solucionó el error de CORS en producción causado por apuntar el backend al dominio viejo del frontend.",
        "Se corrigió el comportamiento de partículas del fondo al cambiar tamaño de ventana y su impacto visual en pantallas pequeñas.",
        "Se restauró la estética correcta del bloque visual de estado para los ESP32 y la interacción para desvincularlos.",
    ]
    for item in items:
        add_bullet(doc, item)


def add_frontend_section(doc):
    doc.add_heading("3. Mejoras de frontend y experiencia de usuario", level=1)
    doc.add_paragraph("Las mejoras visuales y funcionales más relevantes fueron las siguientes:")
    items = [
        "Integración de ThemeProvider y ThemeSelector para aplicar temas reales, tamaño de texto y modo compacto.",
        "Integración de MaintenanceToggle dentro de PlantCard con actualización visual inmediata.",
        "Mejora de ProfileModal, SettingsModal y WelcomeToast para una experiencia más limpia y presentable.",
        "Rediseño de SectorPage y PlantCard con una apariencia más profesional para exposición y mejor jerarquía visual.",
        "Limpieza de textos visibles, iconos y consistencia tipográfica en tarjetas, dashboard y páginas de sector.",
        "Mejora del comportamiento del menú del usuario y del panel de dispositivos ESP32 para vincular y desvincular desde la interfaz.",
        "Ajustes para que la página se sienta más nativa en móvil, con mejores áreas táctiles y menos dependencia de hover.",
    ]
    for item in items:
        add_bullet(doc, item)


def add_perf_section(doc):
    doc.add_heading("4. Optimización de rendimiento", level=1)
    doc.add_paragraph(
        "Se hizo una pasada intensiva de optimización tanto en desktop como en móvil para reducir lag, mejorar la fluidez general "
        "y preparar la aplicación para el escenario de 10 plantas distribuidas en 2 sectores."
    )
    items = [
        "Lazy loading de páginas, modales y componentes pesados para reducir el peso del arranque.",
        "Uso de memoización y diferido de búsqueda en componentes donde había renders repetidos.",
        "Reducción de blur, sombras pesadas y transiciones globales costosas para mejorar INP.",
        "Optimización del fondo de partículas, del efecto de lluvia y de animaciones continuas.",
        "Mejora de carga de imágenes con lazy loading, srcSet y tamaños más adecuados.",
        "Reducción del bundle principal desde aproximadamente 279 kB gzip hasta alrededor de 141 kB gzip.",
        "Ajustes en dashboard y socket para evitar polling innecesario y trabajo extra cuando la pestaña está oculta.",
    ]
    for item in items:
        add_bullet(doc, item)


def add_backend_section(doc):
    doc.add_heading("5. Backend, seguridad y separación por usuario", level=1)
    doc.add_paragraph(
        "El backend recibió cambios importantes para pasar de una demo funcional a una base mucho más sólida, especialmente en seguridad y consistencia multiusuario."
    )
    items = [
        "Protección de rutas administrativas y operativas como devices y mqtt, evitando exposición pública innecesaria.",
        "Corrección del bug que permitía introducir lecturas inválidas de humedad en reportes IoT.",
        "Separación de plantas y dispositivos por usuario, evitando que cuentas distintas vieran o controlaran recursos ajenos.",
        "Aislamiento de eventos en socket.io mediante salas por usuario para que los heartbeats y updates no se crucen entre cuentas.",
        "Validación más fuerte de plantas, umbrales de humedad y combinaciones de sector/válvula.",
        "Mejora del scheduler para que el riego programado también publique comandos al hardware real vía MQTT.",
        "Ajuste de cookies de refresh token para producción cross-site en Render, usando sameSite none y secure.",
        "Incorporación de una Content Security Policy real en el servidor en lugar de desactivarla por completo.",
        "Corrección de CORS para aceptar únicamente el frontend nuevo, más orígenes configurados por entorno.",
    ]
    for item in items:
        add_bullet(doc, item)


def add_iot_section(doc):
    doc.add_heading("6. Integración IoT y firmware", level=1)
    doc.add_paragraph(
        "También se trabajó sobre la integración con los ESP32 para que la parte física estuviera mejor alineada con la web."
    )
    items = [
        "Se completó el firmware del sector superior usando la misma base estable del sector inferior.",
        "Se adaptó el backend para aceptar lecturas por plantId y también por valve/valveNumber.",
        "Se mejoró la vinculación y desvinculación de ESP32 desde el frontend para cuentas distintas.",
        "Se agregó una cola offline persistente en firmware usando Preferences para reenviar lecturas perdidas al reconectar.",
        "Se mantuvo la posibilidad de usar firmware con credenciales directas dentro del .ino para carga manual en los ESP32.",
        "Se dejó más clara la relación entre dispositivo, sector, válvula y planta para pruebas reales.",
    ]
    for item in items:
        add_bullet(doc, item)


def add_render_section(doc):
    doc.add_heading("7. Preparación para producción en Render", level=1)
    items = [
        "Alineación de variables de entorno de backend y frontend con los nombres reales usados por el proyecto.",
        "Ajuste de REACT_APP_API_URL al backend final en producción.",
        "Corrección de recursos rotos del manifest, service worker e iconos faltantes.",
        "Corrección de dominios antiguos que seguían apuntando al frontend viejo.",
        "Configuración de rutas SPA para que la aplicación no falle al recargar en login ni en páginas internas.",
        "Validación de despliegue con frontend y backend ya funcionando correctamente en Render.",
    ]
    for item in items:
        add_bullet(doc, item)


def add_code_map(doc):
    doc.add_heading("8. Archivos clave para explicar al profesor", level=1)
    doc.add_paragraph("Estos archivos son los más útiles para mostrar el trabajo realizado directamente en el código:")
    rows = [
        ("frontend/src/components/layout/Navbar.js", "Corrección de dropdowns, portales, menú de usuario y notificaciones"),
        ("frontend/src/components/layout/GlobalSearch.jsx", "Búsqueda global, overlay y mejoras de UX"),
        ("frontend/src/pages/SectorPage.jsx", "Diseño de sectores, filtros, fluidez y visual premium"),
        ("frontend/src/components/plant/PlantCard.js", "Tarjeta principal de planta, acciones, historial y estado"),
        ("frontend/src/pages/Dashboard.js", "Carga principal, optimización general y lectura del sistema"),
        ("frontend/src/components/dashboard/SystemStatus.js", "Estado de ESP32, vinculación y desvinculación"),
        ("frontend/src/styles.css", "Optimización visual, mobile polish y reducción de carga visual"),
        ("backend/server.js", "CORS, seguridad base, Helmet, CSP y arranque del sistema"),
        ("backend/src/routes/iot.routes.js", "Reportes IoT, heartbeats, válvulas y separación por usuario"),
        ("backend/src/jobs/scheduleRunner.js", "Lógica de riego programado y publicación al hardware"),
        ("backend/src/mqtt/mqttClient.js", "Conexión MQTT, heartbeats y procesamiento de lecturas"),
        ("backend/src/routes/auth.routes.js", "Login, refresh token, cookies y recuperación de contraseña"),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Archivo"
    table.rows[0].cells[1].text = "Qué explica"
    set_cell_shading(table.rows[0].cells[0], "D9F2E3")
    set_cell_shading(table.rows[0].cells[1], "D9F2E3")
    for row_data in rows:
        row = table.add_row().cells
        row[0].text = row_data[0]
        row[1].text = row_data[1]
        for cell in row:
            set_cell_margins(cell)


def add_closure(doc):
    doc.add_heading("9. Conclusión y estado actual", level=1)
    doc.add_paragraph(
        "Al cierre de este avance, el sistema ya cuenta con una base funcional y visual mucho más madura. "
        "La aplicación web quedó más estable, más rápida, mejor organizada por usuario, mejor preparada para producción y lista para pruebas de campo con los ESP32."
    )
    doc.add_paragraph(
        "El siguiente paso natural ya no es una ronda grande de rediseño, sino validar en campo el comportamiento con dispositivos reales, "
        "lecturas de humedad, riego manual, riego programado y confirmación de los dos sectores trabajando sobre la misma base de datos en Atlas."
    )
    doc.add_paragraph("")
    doc.add_paragraph("Documento elaborado como apoyo para revisión de código, exposición y avance técnico del proyecto.")


def main():
    doc = Document()
    style_document(doc)
    add_cover(doc)
    add_status_table(doc)
    add_bug_section(doc)
    add_frontend_section(doc)
    add_perf_section(doc)
    add_backend_section(doc)
    add_iot_section(doc)
    add_render_section(doc)
    add_code_map(doc)
    add_closure(doc)
    doc.save(OUTPUT_PATH)


if __name__ == "__main__":
    main()
